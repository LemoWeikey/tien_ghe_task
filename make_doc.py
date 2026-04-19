"""
Tạo file Word hướng dẫn sử dụng dashboard Vietravel Analytics.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("/Users/jamesgatsby/tien_ghe_task/Huong_dan_Dashboard.docx")

doc = Document()

# ---- Style mặc định ----
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def H(text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def P(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def bullet(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table2(rows, headers=None):
    """Bảng 2 cột đơn giản"""
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = "Light Grid Accent 1"
    if headers:
        hdr = tbl.rows[0].cells
        hdr[0].text = headers[0]
        hdr[1].text = headers[1]
        for c in hdr:
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True
    for i, (a, b) in enumerate(rows, start=1):
        tbl.rows[i].cells[0].text = str(a)
        tbl.rows[i].cells[1].text = str(b)


# ================= TRANG BÌA =================
title = doc.add_heading("Hướng dẫn sử dụng", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Vietravel Analytics Dashboard")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Dashboard phân tích hành vi khách hàng & tỷ lệ chốt tour\n").italic = True
p.add_run("Build với Streamlit + Plotly").italic = True

doc.add_paragraph()

# ================= 1. GIỚI THIỆU =================
H("1. Giới thiệu", 1)
P(
    "Dashboard này giúp đội phân tích của Vietravel theo dõi hành vi khách hàng "
    "trên website travel.com.vn, từ lúc khách bắt đầu vào trang cho đến lúc chốt "
    "tour. Dashboard trả lời các câu hỏi chính:"
)
bullet([
    "Khách đến từ kênh nào (UTM source)? Kênh nào đem lại khách chất lượng?",
    "Khách đi qua những trang nào trước khi Add To Cart / Order Booking?",
    "Tỷ lệ chốt tour thay đổi ra sao khi khách xem tour nhiều lần / dành nhiều thời gian?",
    "Thiết bị, khung giờ, điểm đến nào đang có hiệu suất cao nhất?",
])

# ================= 2. CÁCH CHẠY =================
H("2. Cách khởi chạy", 1)
P("Ở terminal, chạy 2 lệnh:")
p = doc.add_paragraph(style="Intense Quote")
p.add_run("cd /Users/jamesgatsby/tien_ghe_task\nstreamlit run dashboard.py").font.name = "Consolas"
P(
    "Browser sẽ tự mở tab http://localhost:8501. Nếu không, mở thủ công. "
    "Muốn dừng server, nhấn Ctrl+C trên terminal."
)

P("Yêu cầu thư viện:")
bullet([
    "streamlit ≥ 1.40",
    "plotly ≥ 6.0",
    "pandas, numpy",
])

# ================= 3. CẤU TRÚC DASHBOARD =================
H("3. Cấu trúc dashboard", 1)
P("Dashboard gồm 3 khối chính:")
bullet([
    "Sidebar bên trái — chứa các bộ lọc (filter) áp dụng cho toàn dashboard.",
    "Hàng KPI đầu trang — 6 chỉ số tổng quan cập nhật theo filter.",
    "6 tab nội dung — mỗi tab trả lời một nhóm câu hỏi phân tích.",
])

# ================= 4. SIDEBAR FILTERS =================
H("4. Sidebar — Bộ lọc", 1)
P(
    "Mọi filter ở sidebar áp dụng đồng thời cho cả 6 tab. Khi đổi filter, các "
    "biểu đồ và KPI tự refresh."
)

table2([
    ("Thiết bị",
     "Lọc theo loại thiết bị: mobile / pc / tablet. Dùng để so sánh "
     "hành vi giữa các thiết bị."),
    ("UTM source",
     "Lọc theo nguồn traffic quảng cáo: DIGI_Google, DIGI_Facebook, "
     "zalo, DIGI_Admicro… Giá trị NaN = khách truy cập trực tiếp."),
    ("Quốc gia",
     "Lọc theo mã quốc gia ISO (VN, US, KR…) phát hiện qua geo-IP."),
    ("Loại homepage khỏi duration",
     "Checkbox: khi bật, các URL trang chủ (travel.com.vn/) không "
     "tham gia vào tính duration và aggregate. Mặc định: BẬT."),
    ("Khoảng ngày",
     "Date picker 2 đầu. Chọn khoảng ngày (giờ VN, GMT+7) muốn phân tích."),
], headers=["Filter", "Ý nghĩa"])

P(
    "Phía dưới các filter, sidebar hiển thị tổng số events và sessions còn lại "
    "sau khi lọc, để bạn biết mẫu dữ liệu có đủ lớn không."
)

# ================= 5. KPI TOP =================
H("5. Hàng KPI (6 chỉ số tổng quan)", 1)
table2([
    ("👥 Sessions",
     "Số phiên truy cập (session_id duy nhất). Mỗi lần khách mở web "
     "một đợt liên tục = 1 session."),
    ("🙋 Unique users",
     "Số profile khác nhau (profile_id). Một user có thể có nhiều "
     "session, nên con số này thường gần bằng Sessions nhưng nhỏ hơn."),
    ("📊 Events",
     "Tổng số events (mọi hành động: Page View, Search Tour, "
     "Add To Cart…) trong khoảng lọc."),
    ("🛒 Add To Cart",
     "Số session có ít nhất 1 event Add To Cart. Badge % phía dưới = "
     "tỷ lệ trên tổng sessions."),
    ("📝 Order Booking",
     "Số session có event Order Booking (khách điền form đặt tour). "
     "Badge % = tỷ lệ trên tổng sessions."),
    ("✅ Booking Success",
     "Số session có event Booking Success (chốt thành công, có mã "
     "booking). Đây là conversion thực sự."),
], headers=["KPI", "Ý nghĩa"])

# ================= 6. TAB 1: FUNNEL =================
doc.add_page_break()
H("6. Tab 🎯 Funnel", 1)
P(
    "Trả lời câu hỏi: Khách rơi rụng ở bước nào trong quy trình mua tour?"
)

H("6.1. Funnel Chart (biểu đồ phễu)", 2)
P(
    "Biểu đồ phễu 5 bước xếp dọc, chiều rộng mỗi bậc tỉ lệ với số session đến "
    "được bước đó:"
)
bullet([
    "Page View — đã xem ít nhất 1 trang",
    "Search Tour — đã tìm tour",
    "Add To Cart — đã thêm tour vào giỏ",
    "Order Booking — đã điền form đặt tour",
    "Booking Success — đã chốt thành công",
])
P(
    "Trong mỗi thanh phễu hiển thị: giá trị tuyệt đối (số session) + % so với "
    "bước liền trước. Di chuột vào thanh để xem chi tiết."
)
P("Cách đọc:", bold=True)
bullet([
    "Drop-off lớn giữa 2 bước = điểm nghẽn cần tối ưu UX/UI.",
    "Với data hiện tại: Search Tour → Add To Cart mất 95% khách — đây là "
    "điểm nghẽn nghiêm trọng nhất.",
])

H("6.2. Bảng funnel (bên phải)", 2)
P("Cùng số liệu nhưng dạng bảng, tiện copy sang slide/báo cáo.")

H("6.3. Funnel theo thiết bị", 2)
P(
    "Grouped bar chart: mỗi bước trong funnel có 3 thanh cạnh nhau (mobile/pc/"
    "tablet), cho thấy conversion rate từng device. Cách đọc:"
)
bullet([
    "Nếu thanh mobile cao hơn pc ở bước Add To Cart → mobile chốt tốt hơn.",
    "Nếu tablet drop rất sớm → có thể UI tablet có lỗi.",
])

# ================= 7. TAB 2: JOURNEY =================
H("7. Tab 🧭 Customer Journey", 1)
P(
    "Trả lời câu hỏi: Khách đi qua những trang/hành động nào trước khi chốt "
    "(hoặc bỏ đi)?"
)

H("7.1. Top 15 event patterns", 2)
P(
    "Horizontal bar chart hiển thị 15 'đường đi' phổ biến nhất (đã collapse "
    "event trùng liên tiếp). Ví dụ path 'Page View > Search Tour > "
    "Add To Cart'. Text label hiển thị % trên tổng sessions."
)
P("Cách đọc:", bold=True)
bullet([
    "Path chỉ có 1 event (như 'Page View') = bounce — khách xem rồi rời đi.",
    "Path dài, có bước Add To Cart = journey chuyển đổi thành công.",
    "Hiếm khi có path 'Search Tour > Page View > Add To Cart' → khách "
    "tìm, click vào tour, chốt. Đây là journey ideal.",
])

H("7.2. Bảng 30 pattern (bên phải)", 2)
P("Hiển thị đầy đủ 30 pattern hàng đầu, có kèm số sessions và share %.")

H("7.3. Session Explorer", 2)
P(
    "Công cụ lọc session chi tiết. Có 3 filter phụ:"
)
table2([
    ("Min events / session",
     "Chỉ hiện session có ≥ N events. Mặc định = 3 để bỏ qua bounce "
     "1-event (chiếm 67% data) vốn không có journey."),
    ("Chỉ session có Add To Cart / Order",
     "Checkbox: bật để chỉ xem các session đã chốt ý định mua."),
    ("Chỉ session có Booking Success",
     "Checkbox: bật để chỉ xem 5 session đã chốt thành công."),
], headers=["Điều khiển", "Tác dụng"])
P(
    "Bên dưới là bảng tối đa 200 session khớp filter, sort theo num_events "
    "giảm dần. Cột 'had_*' = True/False thể hiện session đó có đạt tới bước "
    "tương ứng không."
)

H("7.4. Dropdown 'Chọn session để xem chi tiết'", 2)
P(
    "Chọn 1 session_id từ danh sách trên → hiển thị timeline chi tiết từng "
    "event của session đó:"
)
bullet([
    "create_ts — thời điểm event xảy ra",
    "name/type — loại event",
    "link — URL trang",
    "context_page_ld_name — tên tour (nếu là page tour)",
    "duration_sec — thời gian ở trang đó (giây)",
])
P(
    "Dùng để kể 'câu chuyện' của 1 khách cụ thể khi trình bày insight hoặc "
    "debug hành vi bất thường.", italic=True
)

# ================= 8. TAB 3: TOUR CONVERSION =================
doc.add_page_break()
H("8. Tab 💼 Tour Conversion", 1)
P(
    "Tab QUAN TRỌNG NHẤT — trả lời câu hỏi gốc: "
    "Tỷ lệ chốt tour có liên quan tới số lần khách xem và thời gian họ ở trang tour không?"
)
P(
    "Phép đo:",
    bold=True,
)
bullet([
    "Unit of analysis = (profile_id, tour_pid): mỗi cặp 1 dòng.",
    "views = số Page View của user trên detail page của tour đó.",
    "total_time = tổng duration_sec của các view đó (đã lọc 0 < d ≤ 30 phút).",
    "intent = True nếu (user, tour) có ít nhất 1 Add To Cart / Order Booking.",
    "booked = True nếu (user, tour) có Order Booking.",
])

H("8.1. Biểu đồ cột 'Intent rate theo số VIEWS'", 2)
P(
    "Trục X là bucket số lần xem tour (1 / 2 / 3-4 / 5-9 / 10+). Trục Y là "
    "intent rate % trong từng bucket. Hover để xem thêm số users và booked rate."
)
P("Cách đọc:", bold=True)
bullet([
    "Nếu bar càng cao khi views tăng → 'xem nhiều → chốt nhiều' (correlation dương).",
    "Data hiện tại: 1-view → 0.83%, 2-view → 3.33% (cao gấp ~4 lần). Xem ≥2 "
    "lần là tín hiệu intent tốt.",
    "Lưu ý: bucket 3-4, 5-9, 10+ có mẫu quá nhỏ nên kém tin cậy.",
])

H("8.2. Biểu đồ cột 'Intent rate theo TỔNG THỜI GIAN'", 2)
P(
    "Giống 8.1 nhưng X là bucket tổng thời gian (<30s, 30-60s, 1-3min, "
    "3-10min, 10min+). Bucket '0 (no duration)' = event cuối session không "
    "đo được."
)
P("Cách đọc:", bold=True)
bullet([
    "Quan hệ không đơn điệu: <30s có intent cao nhất (6.67%) — khách biết "
    "mình muốn gì, bấm ngay.",
    "10min+ có intent thấp (1.37%) — xem lâu có thể là đang phân vân, chưa chắc chốt.",
])

H("8.3. Heatmap 'Views × Time → Intent rate'", 2)
P(
    "Ma trận 2D: hàng = bucket số views, cột = bucket thời gian, màu ô = "
    "intent rate %. Màu càng xanh đậm → tỷ lệ càng cao."
)
P("Cách dùng:", bold=True)
bullet([
    "Tìm ô xanh đậm nhất → định nghĩa 'segment high-intent': VD 'user xem "
    "≥2 lần và dành 1-3 phút'.",
    "Ô trống/xám = không có user ở bucket đó → mẫu quá thưa.",
])

H("8.4. Bảng Correlation", 2)
P("Hai loại correlation giữa các feature và intent (0/1):")
bullet([
    "pearson_r — tương quan tuyến tính. Giá trị gần 0 = không liên quan tuyến tính.",
    "spearman_r — tương quan hạng. Nhạy với quan hệ đơn điệu không tuyến tính.",
    "|r| < 0.1 → yếu; 0.1-0.3 → trung bình; > 0.3 → mạnh.",
])
P(
    "Data hiện tại: r ≈ 0.02-0.10 → correlation YẾU. Kết luận: views/time "
    "một mình không đủ predict chốt tour; cần thêm feature khác (giá, "
    "mùa, kênh).",
    italic=True,
)

H("8.5. Bảng 'Top tour thu hút nhất'", 2)
P("Top 20 tour có nhiều views nhất, kèm:")
bullet([
    "total_views, unique_users, total_time — mức độ quan tâm",
    "intents, bookings — số user đã có ý định / đã chốt",
    "conv_% = intents / unique_users — tỷ lệ chuyển đổi của tour đó",
])
P(
    "Ứng dụng: xác định tour 'hot traffic nhưng chuyển đổi kém' để điều tra "
    "(giá, nội dung trang), và tour 'conv cao' để đẩy ngân sách quảng cáo.",
    italic=True,
)

# ================= 9. TAB 4: DESTINATIONS & UTM =================
doc.add_page_break()
H("9. Tab 🌏 Destinations & UTM", 1)

H("9.1. Top điểm đến", 2)
P(
    "Horizontal bar chart top 15 destination (trích từ URL pattern "
    "/du-lich-<destination>.aspx). Text = số views, hover để xem sessions/users."
)
P("Cách đọc:", bold=True)
bullet([
    "Xếp hạng nhu cầu theo điểm đến: Trung Quốc, Thái Lan, Nhật Bản… dẫn đầu.",
    "Kết hợp với Tab 3 để xem điểm đến nào có conv rate cao → ưu tiên budget ads.",
])

H("9.2. Scatter plot 'UTM Source performance'", 2)
P(
    "Scatter bubble chart: X = số sessions (log scale), Y = ATC rate %, size "
    "bubble = sessions, màu = utm_source. Text label ngay trên bubble."
)
P("Cách đọc:", bold=True)
bullet([
    "Góc phải-trên = 'nhiều traffic + conv cao' → kênh hiệu quả nhất.",
    "Góc phải-dưới = 'nhiều traffic nhưng conv thấp' → đốt tiền, cần xem lại.",
    "Góc trái-trên = 'traffic ít nhưng chất lượng' → đáng scale up.",
    "Log scale trục X giúp thấy rõ cả kênh nhỏ lẫn kênh lớn.",
])
P("Bên dưới là bảng UTM đầy đủ với 3 cột: sessions, atc_pct, order_pct.")

# ================= 10. TAB 5: DEVICE =================
H("10. Tab 📱 Device", 1)

H("10.1. Donut chart 'Tỷ lệ sessions theo thiết bị'", 2)
P("Tỉ trọng mobile/pc/tablet trong tổng sessions.")

H("10.2. Grouped bar 'Conversion rate theo device'", 2)
P(
    "Mỗi device có 2 thanh: atc_pct và order_pct, so sánh tỷ lệ chốt giữa "
    "các loại thiết bị."
)
P("Insight điển hình:", bold=True)
bullet([
    "Mobile dominate traffic (64%) VÀ có order rate cao nhất → checkout mobile "
    "là kênh chốt chính, phải tối ưu UX mobile trước.",
    "PC có share 35% nhưng order rate = 0 → có thể flow desktop bị lỗi.",
])

H("10.3. Bảng device đầy đủ", 2)
P("sessions, share_%, atc_pct, order_pct, avg_events, avg_duration_min.")

H("10.4. OS breakdown", 2)
P(
    "Horizontal bar chart top OS theo số sessions. Để xem phân bổ iOS/Android/"
    "Windows/macOS — hữu ích cho test tương thích."
)

# ================= 11. TAB 6: TIME =================
H("11. Tab ⏰ Time patterns", 1)

H("11.1. Dual-axis chart theo giờ", 2)
P(
    "Hai trục Y trên cùng 1 chart:"
)
bullet([
    "Trục trái (bar xanh) — Page Views theo giờ (0-23h, GMT+7).",
    "Trục phải (line đỏ) — Intent events (ATC + Order) theo giờ.",
])
P("Cách đọc:", bold=True)
bullet([
    "Cao điểm traffic ≠ cao điểm chốt → identify khung giờ 'chất lượng'.",
    "Ví dụ: 9-11h trưa và 20-22h tối thường là 'decision hour' khi khách "
    "ngồi thong thả tìm tour.",
    "Dùng để set lịch ads, lịch chat support, lịch push notification.",
])

H("11.2. Line chart 'Xu hướng theo ngày'", 2)
P(
    "Multi-line: mỗi bước funnel 1 đường (Page View, Search, ATC, Order, "
    "Success) theo ngày. Trục Y dùng log scale để thấy cả bar lớn lẫn bé."
)
P("Cách đọc:", bold=True)
bullet([
    "Spike ngày cụ thể = campaign hoặc sự kiện → đối chiếu với kế hoạch marketing.",
    "Các đường đi song song = funnel ổn định. Đường ATC giảm mà Page View "
    "vẫn tăng = conv rate đang giảm, cần cảnh báo.",
])

# ================= 12. QUY TRÌNH ĐỀ XUẤT =================
doc.add_page_break()
H("12. Quy trình phân tích đề xuất", 1)
P("Sử dụng dashboard theo trình tự sau cho mỗi lần review định kỳ:")

steps = [
    ("B1", "Xem KPI top + Tab Funnel", "Nắm tổng quan sức khỏe funnel hiện tại."),
    ("B2", "Tab Time patterns", "Có spike bất thường không? Ngày nào conv giảm?"),
    ("B3", "Tab Device", "Thiết bị nào đang gánh phần lớn chốt? Có device bị hỏng flow không?"),
    ("B4", "Tab Destinations & UTM", "Kênh nào + điểm đến nào đang hiệu quả? Cần shift budget không?"),
    ("B5", "Tab Tour Conversion", "Tour nào 'hot traffic nhưng conv kém' cần fix landing page?"),
    ("B6", "Tab Journey + Session Explorer", "Deep dive 1-2 session Booking Success để hiểu 'happy path'."),
]
tbl = doc.add_table(rows=1 + len(steps), cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "Bước"
hdr[1].text = "Làm gì"
hdr[2].text = "Mục tiêu"
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
for i, (a, b, c) in enumerate(steps, start=1):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
    tbl.rows[i].cells[2].text = c

# ================= 13. GIỚI HẠN DỮ LIỆU =================
H("13. Các giới hạn cần lưu ý", 1)
bullet([
    "67% session chỉ có 1 event → bounce; không có journey và không đo được duration.",
    "Event cuối cùng của mỗi session không có 'next event' → duration = NaN, loại.",
    "Duration bị cap ở 30 phút để loại user bỏ tab / đi vắng.",
    "Booking Success chỉ có 5 mẫu → thống kê kém tin cậy, nên dùng Add To Cart làm proxy intent.",
    "Sample nhỏ ở các bucket cao (3-4, 5-9, 10+ views) → kết luận cần cautious.",
    "Thời gian là UTC trong data gốc, đã convert +7h sang giờ VN để hiển thị.",
])

# ================= 14. FILE LIÊN QUAN =================
H("14. Các file liên quan trong thư mục", 1)
table2([
    ("full_data.csv", "Dữ liệu thô 14,536 rows × 179 cột (tiếng Việt bị lỗi encoding)."),
    ("full_data_clean.csv", "Dữ liệu đã fix encoding + drop cột HTTP/metadata, còn 67 cột."),
    ("full_data_event_duration.csv", "Bảng duration từng event (1,615 rows có duration hợp lệ)."),
    ("full_data_link_duration.csv", "Aggregate duration theo link (views/avg/median/p90)."),
    ("analysis_report.xlsx", "Báo cáo Excel 14 sheet: funnel, patterns, conversion, correlation…"),
    ("dashboard.py", "Source code Streamlit dashboard."),
    ("process_full.py", "Script sinh ra các file clean + duration."),
    ("analysis_journey.py", "Script sinh analysis_report.xlsx."),
], headers=["File", "Nội dung"])

# ================= KẾT =================
doc.add_page_break()
P("— Hết —", bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUT)
print(f"Đã tạo: {OUT}")
